import subprocess
import sys
import os

target_dir = "C:\\Users\\杢之助\\2nd-Brain\\Clippings\\tmp_env"
if not os.path.exists(target_dir):
    os.makedirs(target_dir)
subprocess.run([sys.executable, "-m", "pip", "install", "docx2txt", "python-docx", "-t", target_dir], check=True, capture_output=True)
sys.path.append(target_dir)

try:
    import docx
    doc = docx.Document("C:\\Users\\杢之助\\2nd-Brain\\Clippings\\260330_水技C　懸案等整理票③（人事）.docx")
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text.strip():
                    row_data.append(cell.text.strip())
            if row_data:
                text.append(" | ".join(row_data))
    print("\n".join(text))
except Exception as e:
    import docx2txt
    print("Using docx2txt instead due to error:", e)
    print(docx2txt.process("C:\\Users\\杢之助\\2nd-Brain\\Clippings\\260330_水技C　懸案等整理票③（人事）.docx"))
